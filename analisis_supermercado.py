import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches

# Configuración de estilo para visualizaciones
sns.set_style("whitegrid")
plt.rcParams['figure.figsize'] = (10, 6)

def cargar_datos(ruta_archivo):
    """Carga el archivo CSV con configuraciones adecuadas para el formato del dataset"""
    try:
        df = pd.read_csv(ruta_archivo, sep=';', decimal=',')
        print(f"Archivo cargado exitosamente. Filas: {len(df)}, Columnas: {len(df.columns)}")
        return df
    except FileNotFoundError:
        print(f"Error: No se encontró el archivo {ruta_archivo}")
        exit()

def explorar_datos(df):
    """Muestra las primeras filas, información general, estadísticas y valores nulos"""
    print("\n--- 1. Carga y Exploración de Datos ---")
    print("\nPrimeras 5 filas del dataset:")
    print(df.head())
    print("\nInformación general del dataset (tipos de datos, no nulos):")
    print(df.info())
    print("\nEstadísticas descriptivas:")
    print(df.describe(include='all'))
    print("\nConteo de valores nulos por columna:")
    print(df.isnull().sum())

def limpiar_datos(df):
    """Maneja valores nulos, convierte fechas y asegura tipos de datos correctos"""
    print("\n--- 2. Limpieza de Datos ---")
    
    # Convertir columnas de fecha al formato datetime
    print("Convirtiendo columnas de fecha (Order Date, Ship Date) a datetime...")
    df['Order Date'] = pd.to_datetime(df['Order Date'], format='%d/%m/%Y')
    df['Ship Date'] = pd.to_datetime(df['Ship Date'], format='%d/%m/%Y')
    print("Fechas convertidas correctamente.")
    
    # Verificar y manejar valores nulos
    nulos = df.isnull().sum()
    if nulos.sum() > 0:
        print("Valores nulos detectados. Procediendo a manejarlos:")
        for col in df.columns:
            if nulos[col] > 0:
                if pd.api.types.is_numeric_dtype(df[col]):
                    mediana = df[col].median()
                    df[col] = df[col].fillna(mediana)
                    print(f"Columna {col}: {nulos[col]} valores nulos rellenados con la mediana ({mediana:.2f})")
                else:
                    moda = df[col].mode()[0] if not df[col].mode().empty else 'Desconocido'
                    df[col] = df[col].fillna(moda)
                    print(f"Columna {col}: {nulos[col]} valores nulos rellenados con la moda ({moda})")
    else:
        print("No se encontraron valores nulos en el dataset.")
    
    # Asegurar tipos de datos numéricos correctos
    columnas_numericas = ['Discount', 'Unit Price', 'Shipping Cost', 'Product Base Margin', 
                          'Profit', 'Quantity ordered new', 'Sales']
    for col in columnas_numericas:
        if col in df.columns:
            df[col] = pd.to_numeric(df[col], errors='coerce')
    
    print("Limpieza de datos completada.")
    return df

def ejecutar_analisis(df):
    """Realiza todos los análisis requeridos y devuelve los resultados"""
    print("\n--- 3. Análisis de Datos ---")
    resultados = {}
    
    # a) Promedio de ventas por categoría
    print("\n3.1 Promedio de ventas por categoría:")
    promedio_ventas_cat = df.groupby('Product Category')['Sales'].mean().reset_index()
    promedio_ventas_cat = promedio_ventas_cat.sort_values(by='Sales', ascending=False)
    resultados['promedio_ventas_cat'] = promedio_ventas_cat
    print(promedio_ventas_cat)
    print("\nConclusión:")
    print(f"La categoría con mayor promedio de ventas es {promedio_ventas_cat.iloc[0]['Product Category']} con {promedio_ventas_cat.iloc[0]['Sales']:.2f}.")
    print("Decisión sugerida: Priorizar inventario y promociones para esta categoría.")
    
    # b) Top 5 ciudades con mayor ganancia
    print("\n3.2 Top 5 ciudades con mayor ganancia (Profit):")
    top_ciudades = df.groupby('City')['Profit'].sum().reset_index()
    top_ciudades = top_ciudades.sort_values(by='Profit', ascending=False).head(5)
    resultados['top_ciudades'] = top_ciudades
    print(top_ciudades)
    print("\nConclusión:")
    print(f"Las ciudades líderes en ganancias son: {', '.join(top_ciudades['City'].tolist())}.")
    print("Decisión sugerida: Enfocar esfuerzos de marketing en estas ciudades.")
    
    # c) Top 5 clientes con mayores compras
    print("\n3.3 Top 5 clientes con mayores compras (Ventas totales):")
    top_clientes = df.groupby('Customer Name')['Sales'].sum().reset_index()
    top_clientes = top_clientes.sort_values(by='Sales', ascending=False).head(5)
    resultados['top_clientes'] = top_clientes
    print(top_clientes)
    print("\nConclusión:")
    print(f"Los clientes con mayores compras son: {', '.join(top_clientes['Customer Name'].tolist())}.")
    print("Decisión sugerida: Implementar programa de lealtad para estos clientes.")
    
    # d) Correlación entre descuentos y ganancias
    print("\n3.4 Correlación entre descuentos (Discount) y ganancias (Profit):")
    corr_descuento_ganancia = df['Discount'].corr(df['Profit'])
    resultados['corr_descuento_ganancia'] = corr_descuento_ganancia
    print(f"Coeficiente de correlación: {corr_descuento_ganancia:.4f}")
    print("\nConclusión:")
    if corr_descuento_ganancia < -0.5:
        print("Correlación negativa fuerte: mayores descuentos están asociados con menores ganancias.")
    elif corr_descuento_ganancia < 0:
        print("Correlación negativa moderada/débil entre descuentos y ganancias.")
    else:
        print("Correlación positiva o nula entre descuentos y ganancias.")
    print("Decisión sugerida: Evaluar rentabilidad de los descuentos aplicados.")
    
    # Análisis adicional 1: Ventas por mes
    print("\n--- 4. Análisis Adicional ---")
    print("\n4.1 Ventas totales por mes:")
    df['Month'] = df['Order Date'].dt.month
    df['Month_Name'] = df['Order Date'].dt.strftime('%B')
    ventas_por_mes = df.groupby(['Month', 'Month_Name'])['Sales'].sum().reset_index()
    ventas_por_mes = ventas_por_mes.sort_values(by='Month')
    resultados['ventas_por_mes'] = ventas_por_mes
    print(ventas_por_mes[['Month_Name', 'Sales']])
    print("\nConclusión:")
    mes_max = ventas_por_mes.loc[ventas_por_mes['Sales'].idxmax()]
    print(f"Mes con mayores ventas: {mes_max['Month_Name']}. Mes con menores: {ventas_por_mes.loc[ventas_por_mes['Sales'].idxmin()]['Month_Name']}.")
    print("Decisión sugerida: Planificar promociones en meses de baja venta.")
    
    # Análisis adicional 2: Productos con pérdidas
    print("\n4.2 Productos con pérdidas (Profit < 0):")
    productos_perdida = df[df['Profit'] < 0].groupby('Product Name')['Profit'].sum().reset_index()
    productos_perdida = productos_perdida.sort_values(by='Profit', ascending=True).head(10)
    resultados['productos_perdida'] = productos_perdida
    print(f"Total de registros con pérdidas: {len(df[df['Profit'] < 0])}")
    print("Top 10 productos con mayores pérdidas:")
    print(productos_perdida)
    print("\nConclusión:")
    print(f"Pérdidas totales de los top 10 productos: {productos_perdida['Profit'].sum():.2f}.")
    print("Decisión sugerida: Revisar costos de estos productos o descontinuarlos.")
    
    return resultados, df

def generar_visualizaciones(resultados, df):
    """Crea las gráficas requeridas y las guarda como archivos PNG"""
    print("\n--- 5. Generación de Visualizaciones ---")
    archivos_plot = []
    
    # 1. Barras: Promedio de ventas por categoría
    plt.figure()
    sns.barplot(x='Product Category', y='Sales', data=resultados['promedio_ventas_cat'], hue='Product Category', palette='viridis', legend=False)
    plt.title('Promedio de Ventas por Categoria')
    plt.xlabel('Categoria')
    plt.ylabel('Promedio de Ventas')
    plt.xticks(rotation=45)
    plt.tight_layout()
    ruta = 'promedio_ventas_categoria.png'
    plt.savefig(ruta)
    archivos_plot.append(ruta)
    plt.close()
    print(f"Guardada: {ruta}")
    
    # 2. Top 5 ciudades con mayor ganancia
    plt.figure()
    sns.barplot(x='Profit', y='City', data=resultados['top_ciudades'], hue='City', palette='magma', legend=False)
    plt.title('Top 5 Ciudades con Mayor Ganancia')
    plt.xlabel('Ganancia Total')
    plt.ylabel('Ciudad')
    plt.tight_layout()
    ruta = 'top_ciudades.png'
    plt.savefig(ruta)
    archivos_plot.append(ruta)
    plt.close()
    print(f"Guardada: {ruta}")
    
    # 3. Ventas por mes (línea)
    plt.figure()
    plt.plot(resultados['ventas_por_mes']['Month'], resultados['ventas_por_mes']['Sales'], marker='o', color='b')
    plt.title('Ventas Totales por Mes')
    plt.xlabel('Mes')
    plt.ylabel('Ventas Totales')
    plt.xticks(resultados['ventas_por_mes']['Month'], resultados['ventas_por_mes']['Month_Name'], rotation=45)
    plt.grid(True)
    plt.tight_layout()
    ruta = 'ventas_por_mes.png'
    plt.savefig(ruta)
    archivos_plot.append(ruta)
    plt.close()
    print(f"Guardada: {ruta}")
    
    # 4. Descuento vs Ganancia (scatter plot)
    plt.figure()
    sns.scatterplot(x='Discount', y='Profit', data=df, alpha=0.6, color='g')
    plt.title('Relacion entre Descuento y Ganancia')
    plt.xlabel('Descuento')
    plt.ylabel('Ganancia')
    plt.tight_layout()
    ruta = 'descuento_vs_ganancia.png'
    plt.savefig(ruta)
    archivos_plot.append(ruta)
    plt.close()
    print(f"Guardada: {ruta}")
    
    # 5. Top 10 productos con pérdidas
    plt.figure()
    sns.barplot(x='Profit', y='Product Name', data=resultados['productos_perdida'], hue='Product Name', palette='rocket', legend=False)
    plt.title('Top 10 Productos con Mayores Perdidas')
    plt.xlabel('Perdida Total')
    plt.ylabel('Producto')
    plt.tight_layout()
    ruta = 'productos_perdida.png'
    plt.savefig(ruta)
    archivos_plot.append(ruta)
    plt.close()
    print(f"Guardada: {ruta}")
    
    return archivos_plot

def generar_reporte_word(resultados, archivos_plot, df):
    """Crea el archivo Word con todos los análisis, resultados y visualizaciones"""
    print("\n--- 6. Generación de Reporte Word ---")
    doc = Document()
    
    # Título
    doc.add_heading('Analisis de Datos de Ventas de Supermercado', 0)
    
    # Introducción
    doc.add_heading('Introduccion', 1)
    doc.add_paragraph('Este documento presenta el analisis detallado de los datos de ventas del supermercado correspondientes al año 2015, utilizando el dataset SuperStoreUS-2015(Orders).csv. El analisis incluye limpieza de datos, metricas clave, visualizaciones y conclusiones para la toma de decisiones.')
    doc.add_paragraph(f'El dataset contiene {len(df)} registros y {len(df.columns)} columnas con informacion de pedidos, clientes, productos, ventas y ganancias.')
    
    # Análisis 1: Promedio de ventas por categoría
    doc.add_heading('1. Promedio de Ventas por Categoria', 1)
    doc.add_paragraph('Calcula el promedio de ventas para cada categoria de producto.')
    doc.add_paragraph('Resultados:')
    tabla = doc.add_table(rows=1, cols=2)
    tabla.rows[0].cells[0].text = 'Categoria'
    tabla.rows[0].cells[1].text = 'Promedio de Ventas'
    for _, fila in resultados['promedio_ventas_cat'].iterrows():
        nueva_fila = tabla.add_row()
        nueva_fila.cells[0].text = fila['Product Category']
        nueva_fila.cells[1].text = f"{fila['Sales']:.2f}"
    doc.add_picture(archivos_plot[0], width=Inches(6))
    doc.add_paragraph('Conclusion: La categoria con mayor promedio de ventas debe ser priorizada en inventario y promociones.')
    
    # Análisis 2: Top 5 ciudades con mayor ganancia
    doc.add_heading('2. Top 5 Ciudades con Mayor Ganancia', 1)
    doc.add_paragraph('Identifica las ciudades con mayores ganancias totales.')
    doc.add_paragraph('Resultados (Top 5):')
    tabla = doc.add_table(rows=1, cols=2)
    tabla.rows[0].cells[0].text = 'Ciudad'
    tabla.rows[0].cells[1].text = 'Ganancia Total'
    for _, fila in resultados['top_ciudades'].iterrows():
        nueva_fila = tabla.add_row()
        nueva_fila.cells[0].text = fila['City']
        nueva_fila.cells[1].text = f"{fila['Profit']:.2f}"
    doc.add_picture(archivos_plot[1], width=Inches(6))
    doc.add_paragraph('Conclusion: Enfocar esfuerzos de marketing en estas ciudades para maximizar ganancias.')
    
    # Análisis 3: Top 5 clientes con mayores compras
    doc.add_heading('3. Top 5 Clientes con Mayores Compras', 1)
    doc.add_paragraph('Identifica los clientes con mayor gasto acumulado.')
    doc.add_paragraph('Resultados (Top 5):')
    tabla = doc.add_table(rows=1, cols=2)
    tabla.rows[0].cells[0].text = 'Cliente'
    tabla.rows[0].cells[1].text = 'Ventas Totales'
    for _, fila in resultados['top_clientes'].iterrows():
        nueva_fila = tabla.add_row()
        nueva_fila.cells[0].text = fila['Customer Name']
        nueva_fila.cells[1].text = f"{fila['Sales']:.2f}"
    doc.add_paragraph('Conclusion: Implementar programas de lealtad para retener a estos clientes.')
    
    # Análisis 4: Correlación descuento vs ganancia
    doc.add_heading('4. Correlacion entre Descuentos y Ganancias', 1)
    doc.add_paragraph('Evalua la relacion entre el porcentaje de descuento y la ganancia.')
    doc.add_paragraph(f'Coeficiente de correlacion: {resultados["corr_descuento_ganancia"]:.4f}')
    doc.add_picture(archivos_plot[3], width=Inches(6))
    if resultados['corr_descuento_ganancia'] < -0.5:
        doc.add_paragraph('Conclusion: Correlacion negativa fuerte, evaluar rentabilidad de descuentos.')
    elif resultados['corr_descuento_ganancia'] < 0:
        doc.add_paragraph('Conclusion: Correlacion negativa moderada entre descuentos y ganancias.')
    else:
        doc.add_paragraph('Conclusion: Correlacion positiva o nula entre descuentos y ganancias.')
    
    # Análisis adicional 1: Ventas por mes
    doc.add_heading('5. Ventas Totales por Mes', 1)
    doc.add_paragraph('Identifica patrones estacionales de ventas.')
    doc.add_paragraph('Resultados:')
    tabla = doc.add_table(rows=1, cols=2)
    tabla.rows[0].cells[0].text = 'Mes'
    tabla.rows[0].cells[1].text = 'Ventas Totales'
    for _, fila in resultados['ventas_por_mes'].iterrows():
        nueva_fila = tabla.add_row()
        nueva_fila.cells[0].text = fila['Month_Name']
        nueva_fila.cells[1].text = f"{fila['Sales']:.2f}"
    doc.add_picture(archivos_plot[2], width=Inches(6))
    doc.add_paragraph('Conclusion: Planificar promociones en meses de baja venta.')
    
    # Análisis adicional 2: Productos con pérdidas
    doc.add_heading('6. Productos con Perdidas (Profit < 0)', 1)
    doc.add_paragraph('Identifica productos que generaron perdidas netas.')
    doc.add_paragraph(f'Total de registros con perdidas: {len(df[df["Profit"] < 0])}')
    doc.add_picture(archivos_plot[4], width=Inches(6))
    doc.add_paragraph('Conclusion: Revisar costos de estos productos o descontinuar los menos rentables.')
    
    # Conclusiones generales
    doc.add_heading('Conclusiones Generales', 1)
    doc.add_paragraph('1. Priorizar la categoria con mayor promedio de ventas en inventario y marketing.')
    doc.add_paragraph('2. Enfocar esfuerzos de expansion en las ciudades con mayores ganancias.')
    doc.add_paragraph('3. Implementar programas de fidelizacion para los clientes con mayores compras.')
    doc.add_paragraph('4. Monitorear la correlacion entre descuentos y ganancias para evitar erosion de margenes.')
    doc.add_paragraph('5. Aprovechar patrones estacionales para campanas promocionales.')
    doc.add_paragraph('6. Revisar costos de productos con perdidas recurrentes.')
    
    # Guardar documento
    doc.save('Analisis_Supermercado.docx')
    print("Reporte Word generado exitosamente: Analisis_Supermercado.docx")

if __name__ == "__main__":
    # Ruta del archivo CSV
    ruta_csv = "SuperStoreUS-2015(Orders).csv"
    
    # Ejecutar pasos del análisis
    df = cargar_datos(ruta_csv)
    explorar_datos(df)
    df = limpiar_datos(df)
    resultados, df = ejecutar_analisis(df)
    archivos_plot = generar_visualizaciones(resultados, df)
    generar_reporte_word(resultados, archivos_plot, df)
    
    print("\n--- Proceso Completado ---")
    print("Todos los analisis, visualizaciones y el reporte Word han sido generados exitosamente.")
